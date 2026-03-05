#!/usr/bin/env python3
"""Generate project documentation template in DOCX format - simple format."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def create_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # ── [Emri i Projektit] - centered title ──
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run('Fiks Platform')
    run_title.font.size = Pt(16)
    run_title.font.bold = True

    # ── Përshkrimi i Projektit ──
    p_label = doc.add_paragraph()
    run_label = p_label.add_run('Përshkrimi i Projektit:')
    run_label.font.bold = True
    run_label.font.size = Pt(12)

    p_desc = doc.add_paragraph()
    run_desc = p_desc.add_run(
        'Fiks Platform është një platformë web që mundëson lidhjen midis klientëve '
        'dhe profesionistëve për shërbime të ndryshme. Përdoruesit mund të kërkojnë '
        'shërbime, të komunikojnë me profesionistë, të lënë vlerësime, dhe të '
        'menaxhojnë profilet e tyre. Platforma përfshin gjithashtu një panel '
        'administrativ për menaxhimin e përdoruesve dhe shërbimeve.'
    )
    run_desc.font.size = Pt(12)

    # ── Anëtarët e Ekipit ──
    p_team_label = doc.add_paragraph()
    run_team = p_team_label.add_run('Anëtarët e Ekipit:')
    run_team.font.bold = True
    run_team.font.size = Pt(12)

    # Numri i anëtarëve
    p_num = doc.add_paragraph()
    p_num.style = doc.styles['List Bullet']
    run_num = p_num.add_run('Numri i anëtarëve: 1 student.')
    run_num.font.size = Pt(12)

    # Team table - 6 columns: ID, Emri, Mbiemri, Email, Grupi, Ligjëruesi
    headers = ['ID', 'Emri', 'Mbiemri', 'Email', 'Grupi', 'Ligjëruesi']
    team_rows = [
        ['2324.....', 'Rron', 'Elshani', 'elshanirron@gmail.com', '4', 'Lavdim Menxhiqi'],
    ]

    table = doc.add_table(rows=1 + len(team_rows), cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for i, row_data in enumerate(team_rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer

    # ── Teknologjitë ──
    p_tech_label = doc.add_paragraph()
    run_tech = p_tech_label.add_run('Teknologjitë:')
    run_tech.font.bold = True
    run_tech.font.size = Pt(12)

    techs = [
        ('Backend', 'Node.js + Express'),
        ('Frontend', 'React + Vite'),
        ('Databaza', 'PostgreSQL 15'),
    ]

    for label, value in techs:
        p = doc.add_paragraph()
        p.style = doc.styles['List Bullet']
        run_l = p.add_run(f'{label} : ')
        run_l.font.bold = True
        run_l.font.size = Pt(12)
        run_v = p.add_run(value)
        run_v.font.size = Pt(12)

    # Save
    output_path = os.path.join(os.path.dirname(__file__), 'Fiks_Platform_Projekt.docx')
    doc.save(output_path)
    print(f'✅ Dokumenti u krijua me sukses: {output_path}')
    return output_path


if __name__ == '__main__':
    create_document()
